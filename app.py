import cv2
from pyzbar.pyzbar import decode
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import os
import requests
import qrcode
from PIL import Image, ImageTk

class SkanerIGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Kombajn do kodów qr i kreskowych")
        self.root.geometry("700x700")
        self.root.configure(bg="#f0f4f8")

        # System zakładek
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        self.tab_scan = tk.Frame(self.notebook, bg="#f0f4f8")
        self.tab_gen = tk.Frame(self.notebook, bg="#f0f4f8")

        self.notebook.add(self.tab_scan, text="Dekodowanie (skaner)")
        self.notebook.add(self.tab_gen, text="Nadawanie (generator)")

        self.buduj_zakladke_skanera()
        self.buduj_zakladke_generatora()

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    # ==========================================
    # ZAKŁADKA 1: SKANER (DEKODOWANIE)
    # ==========================================
    def buduj_zakladke_skanera(self):
        self.is_running = False
        self.found_data = None
        self.found_type = None
        self.extra_info = ""

        label = tk.Label(self.tab_scan, text="Odczytuj kody z kamery lub pliku", font=("Segoe UI", 16, "bold"), bg="#f0f4f8", fg="#2c3e50")
        label.pack(pady=10)

        self.result_frame = tk.Frame(self.tab_scan, bg="#ffffff", bd=2, relief="groove")
        self.result_frame.pack(pady=5, padx=20, fill="x")

        self.type_label = tk.Label(self.result_frame, text="Typ kodu: brak", font=("Segoe UI", 12, "bold"), bg="#ffffff", fg="#e67e22")
        self.type_label.pack(pady=5)

        self.data_label = tk.Label(self.result_frame, text="Zeskanowana treść pojawi się tutaj...", font=("Segoe UI", 11), bg="#ffffff", fg="#34495e", wraplength=600, justify="center")
        self.data_label.pack(pady=10)

        self.btn_start = tk.Button(self.tab_scan, text="📷 Uruchom kamerę w jakości hd", font=("Segoe UI", 11, "bold"), bg="#3498db", fg="white", command=self.start_scanner, width=30, bd=0)
        self.btn_start.pack(pady=5)

        self.btn_file = tk.Button(self.tab_scan, text="📁 Wczytaj zdjęcie z dysku", font=("Segoe UI", 11, "bold"), bg="#e67e22", fg="white", command=self.load_from_file_scan, width=30, bd=0)
        self.btn_file.pack(pady=5)

        self.action_frame = tk.Frame(self.tab_scan, bg="#f0f4f8")
        self.action_frame.pack(pady=10)

        self.btn_copy = tk.Button(self.action_frame, text="📋 Kopiuj wynik", font=("Segoe UI", 10, "bold"), bg="#9b59b6", fg="white", command=self.copy_to_clipboard, width=16, state="disabled", bd=0)
        self.btn_copy.grid(row=0, column=0, padx=10)

        self.btn_save = tk.Button(self.action_frame, text="💾 Zapisz do Worda", font=("Segoe UI", 10, "bold"), bg="#2ecc71", fg="white", command=self.save_to_word, width=16, state="disabled", bd=0)
        self.btn_save.grid(row=0, column=1, padx=10)

        self.status_var = tk.StringVar(value="Status: gotowy do pracy")
        self.status_label = tk.Label(self.tab_scan, textvariable=self.status_var, bg="#f0f4f8", font=("Segoe UI", 10, "italic"), fg="#7f8c8d")
        self.status_label.pack(pady=10)

    def start_scanner(self):
        if self.is_running: return
        self.is_running = True
        self.extra_info = ""
        self.btn_start.config(state="disabled")
        self.btn_file.config(state="disabled")
        self.status_var.set("Status: uruchamianie kamery...")
        threading.Thread(target=self.scan_logic, daemon=True).start()

    def load_from_file_scan(self):
        file_path = filedialog.askopenfilename(title="Wybierz zdjęcie", filetypes=(("Pliki graficzne", "*.png;*.jpg;*.jpeg;*.bmp"), ("Wszystkie pliki", "*.*")))
        if file_path:
            self.extra_info = ""
            self.status_var.set("Status: analizowanie pliku...")
            threading.Thread(target=self.process_image_file, args=(file_path,), daemon=True).start()

    def process_image_file(self, file_path):
        img = cv2.imread(file_path)
        if img is not None:
            codes = decode(img)
            if codes:
                self.dekoduj_dane(codes[0])
                self.process_found_code()
            else:
                self.found_data = None
                self.status_var.set("Status: nie wykryto kodu.")
        else:
            self.status_var.set("Status: błąd pliku.")
        self.root.after(0, self.update_gui_after_scan)

    def scan_logic(self):
        cap = cv2.VideoCapture(0)
        cap.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
        cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)
        win_name = "Podglad z kamery (q aby wyjsc)"
        
        while self.is_running:
            ret, frame = cap.read()
            if not ret: break
            codes = decode(frame)
            if codes:
                self.dekoduj_dane(codes[0])
                self.is_running = False
                break
            cv2.imshow(win_name, frame)
            if cv2.waitKey(1) & 0xFF == ord('q') or cv2.getWindowProperty(win_name, cv2.WND_PROP_VISIBLE) < 1:
                self.is_running = False
                break
                
        cap.release()
        cv2.destroyAllWindows()
        if self.found_data: self.process_found_code()
        else: self.root.after(0, self.update_gui_after_scan)

    def dekoduj_dane(self, kod):
        surowe_bajty = kod.data
        self.found_type = kod.type
        # Tarcza chroniąca przed "chińskimi znakami" - wymuszamy poprawne polskie kodowanie
        try:
            self.found_data = surowe_bajty.decode('utf-8')
        except UnicodeDecodeError:
            try:
                self.found_data = surowe_bajty.decode('cp1250') # Awaryjny tryb Windows
            except:
                self.found_data = surowe_bajty.decode('iso-8859-2', errors='replace') # Stary standard polski

    def process_found_code(self):
        if self.found_type in ['EAN13', 'EAN8', 'UPCA']:
            self.status_var.set("Status: przeszukiwanie baz danych produktów...")
            self.analyze_ean(self.found_data)
        elif "QR" in self.found_type:
            self.status_var.set("Status: dekodowanie zawartości qr...")
            self.analyze_qr(self.found_data)
        self.root.after(0, self.update_gui_after_scan)

    def analyze_qr(self, qr_text):
        tekst = qr_text.upper()
        if tekst.startswith("WIFI:"):
            self.extra_info += "Kategoria: sieć Wi-Fi\n\n"
            for p in qr_text[5:].split(';'):
                if p.startswith('S:'): self.extra_info += f"Nazwa sieci: {p[2:]}\n"
                elif p.startswith('P:'): self.extra_info += f"Hasło: {p[2:]}\n"
        elif tekst.startswith("BEGIN:VCARD"):
            self.extra_info += "Kategoria: wizytówka\n\n"
            for line in qr_text.split('\n'):
                if line.upper().startswith("FN:"): self.extra_info += f"Osoba: {line[3:]}\n"
                elif line.upper().startswith("TEL"): self.extra_info += f"Telefon: {line.split(':')[-1]}\n"
        elif tekst.startswith("HTTP"):
            self.extra_info += "Kategoria: link internetowy\n"
        else:
            self.extra_info += "Kategoria: zwykły tekst / inny format\n"

    def analyze_ean(self, ean):
        kraj = "Nieznany / reszta świata"
        if ean.startswith('590'): kraj = "Polska"
        elif ean.startswith(('40', '41', '42', '43', '44')): kraj = "Niemcy"
        self.extra_info += f"Kraj rejestracji: {kraj}\n\n"

        if ean.startswith('978') or ean.startswith('979'):
            try:
                data = requests.get(f"https://www.googleapis.com/books/v1/volumes?q=isbn:{ean}", timeout=5).json()
                if data.get('totalItems', 0) > 0:
                    b = data['items'][0]['volumeInfo']
                    self.extra_info += f"Książka: {b.get('title', 'Brak')} ({', '.join(b.get('authors', []))})\n"
                    return
            except: pass

        znaleziono = False
        try:
            data_food = requests.get(f"https://world.openfoodfacts.org/api/v0/product/{ean}.json", timeout=5).json()
            if data_food.get('status') == 1 and data_food['product'].get('product_name'):
                self.extra_info += f"Produkt: {data_food['product'].get('product_name')}\n"
                znaleziono = True
        except: pass

        if not znaleziono:
            try:
                data_upc = requests.get(f"https://api.upcitemdb.com/prod/trial/lookup?upc={ean}", headers={'Accept': 'application/json'}, timeout=6).json()
                if data_upc.get('code') == 'OK' and len(data_upc.get('items', [])) > 0:
                    self.extra_info += f"Produkt: {data_upc['items'][0].get('title', 'Brak')}\n"
            except:
                self.extra_info += "Brak w darmowych bazach on-line.\n"

    def update_gui_after_scan(self):
        self.btn_start.config(state="normal")
        self.btn_file.config(state="normal")
        if self.found_data:
            self.type_label.config(text=f"Rozpoznano: {self.found_type}")
            self.data_label.config(text=f"Odczyt:\n{self.found_data}\n\nAnaliza:\n{self.extra_info}" if self.extra_info else self.found_data)
            self.status_var.set("Status: zakończono pomyślnie!")
            self.btn_copy.config(state="normal")
            self.btn_save.config(state="normal")
        else:
            if "błąd" not in self.status_var.get().lower():
                self.status_var.set("Status: przerwano.")

    def copy_to_clipboard(self):
        if self.found_data:
            self.root.clipboard_clear()
            self.root.clipboard_append(f"{self.found_data}\n\n{self.extra_info}")
            messagebox.showinfo("Skopiowano", "Skopiowano do schowka.")

    def save_to_word(self):
        if not self.found_data: return
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Dokument Word", "*.docx")])
        if path:
            doc = Document()
            doc.add_heading('Wynik odczytu', 0)
            doc.add_paragraph(self.found_data)
            if self.extra_info:
                doc.add_heading('Analiza:', 1)
                doc.add_paragraph(self.extra_info)
            doc.save(path)
            os.startfile(path)

    # ==========================================
    # ZAKŁADKA 2: GENERATOR (NADAWANIE)
    # ==========================================
    def buduj_zakladke_generatora(self):
        label = tk.Label(self.tab_gen, text="Stwórz własny kod qr", font=("Segoe UI", 16, "bold"), bg="#f0f4f8", fg="#2c3e50")
        label.pack(pady=10)

        desc = tk.Label(self.tab_gen, text="Wpisz dowolny tekst, adres portfela krypto\nlub wczytaj ścieżkę do zdjęcia na komputerze:", bg="#f0f4f8", font=("Segoe UI", 10))
        desc.pack(pady=5)

        self.text_box = tk.Text(self.tab_gen, height=6, width=55, font=("Segoe UI", 11), bg="white", relief="groove", bd=2)
        self.text_box.pack(pady=10, padx=20)

        self.btn_file_gen = tk.Button(self.tab_gen, text="📁 Pobierz ścieżkę do pliku z komputera", font=("Segoe UI", 10, "bold"), bg="#e67e22", fg="white", command=self.choose_file_for_gen, bd=0, padx=10, pady=5)
        self.btn_file_gen.pack(pady=5)

        self.btn_generate = tk.Button(self.tab_gen, text="🛠️ Generuj kod", font=("Segoe UI", 12, "bold"), bg="#2ecc71", fg="white", command=self.generate_qr, bd=0, width=20, height=2)
        self.btn_generate.pack(pady=15)

        self.qr_label = tk.Label(self.tab_gen, bg="#f0f4f8")
        self.qr_label.pack(pady=10)

        self.gen_save_btn = tk.Button(self.tab_gen, text="💾 Zapisz obrazek kodu", font=("Segoe UI", 10, "bold"), bg="#3498db", fg="white", command=self.save_generated_qr, bd=0, width=20, state="disabled")
        self.gen_save_btn.pack(pady=5)
        self.current_qr_img = None

    def choose_file_for_gen(self):
        file_path = filedialog.askopenfilename(title="Wybierz plik")
        if file_path:
            self.text_box.delete("1.0", tk.END)
            self.text_box.insert(tk.END, file_path)

    def generate_qr(self):
        data = self.text_box.get("1.0", tk.END).strip()
        if not data:
            messagebox.showwarning("Pusto", "Musisz wpisać jakiś tekst lub wybrać plik.")
            return

        try:
            qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=8, border=4)
            # Tu biblioteka naturalnie pakuje poprawny tekst polski
            qr.add_data(data)
            qr.make(fit=True)
            self.current_qr_img = qr.make_image(fill_color="black", back_color="white")
            
            img_tk = ImageTk.PhotoImage(self.current_qr_img.resize((200, 200)))
            self.qr_label.config(image=img_tk)
            self.qr_label.image = img_tk 
            self.gen_save_btn.config(state="normal")
            
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się wygenerować: {e}")

    def save_generated_qr(self):
        if self.current_qr_img:
            path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("Plik PNG", "*.png")])
            if path:
                self.current_qr_img.save(path)
                messagebox.showinfo("Zapisano", "Twój nowy kod qr został zapisany!")

    def on_closing(self):
        self.is_running = False
        self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = SkanerIGenerator(root)
    root.mainloop()
